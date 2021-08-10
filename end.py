from selenium.webdriver.firefox.options import Options as FirefoxOptions
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from time import sleep
import timeit as t
import pandas as pd
from selenium import webdriver
from docxtpl import DocxTemplate
import io
from PIL import Image
from docx.shared import Inches

# Простоянные переменные
driver = webdriver.Firefox()
app_login = 'univer@kaznu.kz'
app_password = 'yhxTQRBc2n'
delay = 9


driver.get('https://app.oqylyq.kz/')
sleep(1.7)
driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div/div/div/div[2]/div[2]/div/form/div[1]/div[1]/div/input').send_keys(app_login)
driver.find_element_by_xpath('//*[@id="__layout"]/div/div[1]/div/div/div/div[2]/div[2]/div/form/div[2]/div/div/input').send_keys(app_password)
driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div/div/div/div[2]/div[2]/div/form/div[4]/div/button/div[2]').click()
sleep(3)
i = 2
# for i in range(1, 23):
driver.get(f'https://app.oqylyq.kz/t/assignments?page=2')
sleep(1)

for j in range(1,10):
    dl = 5
    print(j)
    if j == 11:
        nope = 0
        break
    sleep(2)
    exz = str(driver.find_element_by_xpath(f'/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[2]/div/a[{j}]/div[2]/div[1]/div[1]').text)
    print(exz)
    print(f'/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[2]/div/a[{j}]/div[1]/div/div/div/div[2]/div/div')
    c = driver.find_element_by_xpath(f'/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[2]/div/a[{j}]/div[1]/div/div/div/div[2]/div/div').text
    print(c)
    if c == '0':
        # driver.get(f'https://app.oqylyq.kz/t/assignments?page={i}')
        continue
    sleep(1)


    driver.find_element_by_xpath(f'/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[2]/div/a[{j}]/div[2]/div[1]/div[1]').click()


    sleep(2)
    # driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/a').click()
    student_number = 1
    for i in range(130):
        bl = 1
        rp = 0
        pl = 0
        question3 = ""
        question4 = ""
        B4 = ''
        T5 = ''
        T5B1 = 0
        T5B2 = 0
        T5B3 = 0
        T5B4 = 0
        T1B4, T2B4, T3B4, T4B4, T5B4 = 0, 0, 0, 0, 0
        T2B1, T2B2, T2B3, T2B4, T2B5 = 0, 0, 0, 0, 0
        T3B1, T3B2, T3B3, T3B4, T3B5 = 0, 0, 0, 0, 0
        B4 = 0
        while True:
            try:
                driver.find_element_by_xpath(f'/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/a[{student_number}]/div[2]/div').click()
                break
            except:
                rp += 1
                if rp > 11:
                    break
                continue
        if rp > 11:
            break
        while True:
            try:
                driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[3]/div[2]/span[1]/div').click()
                #/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[3]/div[2]/span[1]/div
                break
            except:
                continue


        sleep(2)
        question1 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[2]/div/div/div[2]').screenshot_as_png
        print(question1)
        question1 = io.BytesIO(question1)
        print(question1)
        sleep(1)
        driver.find_element_by_xpath('/html/body/div[1]/div/div/div[2]/div/div/div[1]').click()
        driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[3]/div[2]/span[2]/div').click()
        sleep(2)
        question2 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[2]/div/div/div[2]').screenshot_as_png
        question2 = io.BytesIO(question2)
        driver.find_element_by_xpath('/html/body/div[1]/div/div/div[2]/div/div/div[1]').click()
        sleep(1)
        try:
            driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[3]/div[2]/span[3]/div').click()
            sleep(2)
            question3 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[2]/div/div/div[2]').screenshot_as_png
            sleep(1)
            question3 = io.BytesIO(question3)
            driver.find_element_by_xpath('/html/body/div[1]/div/div/div[2]/div/div/div[1]').click()
        except:
            print('2 bl')
        try:
            driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[3]/div[2]/span[4]/div').click()
            sleep(1)
            question4 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[2]/div/div/div[2]').screenshot_as_png
            question4 = io.BytesIO(question4)
            driver.find_element_by_xpath('/html/body/div[1]/div/div/div[2]/div/div/div[1]').click()
        except:
            bl = 0
        name = driver.find_element_by_xpath(
            f'/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/a[{student_number}]/div[2]/div').text
        # БАЛЛЫ
        try:
            T1B1 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[4]/div[2]/span[1]/div/div[3]/div/span').text
        except:
            continue
        T1B2 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[4]/div[2]/span[2]/div/div[3]/div/span').text
        T1B3 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[4]/div[2]/span[3]/div/div[3]/div/span').text
        try:
            T2B1 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[5]/div[2]/span[1]/div/div[3]/div/span').text
            T2B2 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[5]/div[2]/span[2]/div/div[3]/div/span').text
            T2B3 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[5]/div[2]/span[3]/div/div[3]/div/span').text
        except:
            dl = 4
        try:
            T3B1 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[6]/div[2]/span[1]/div/div[3]/div/span').text
            T3B2 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[6]/div[2]/span[2]/div/div[3]/div/span').text
            T3B3 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[6]/div[2]/span[3]/div/div[3]/div/span').text
        except:
            dl = 4
        try:
            T4B1 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[7]/div[2]/span[1]/div/div[3]/div/span').text
            T4B2 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[7]/div[2]/span[2]/div/div[3]/div/span').text
            T4B3 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[7]/div[2]/span[3]/div/div[3]/div/span').text
        except:
            dl = 4
        try:
            T5B1 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[8]/div[2]/span[1]/div/div[3]/div/span').text
            T5B2 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[8]/div[2]/span[2]/div/div[3]/div/span').text
            T5B3 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[8]/div[2]/span[3]/div/div[3]/div/span').text
        except:
            dl = 4
        if bl:
            T1B4 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[4]/div[2]/span[4]/div/div[3]/div/span').text
            try:
                T2B4 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[5]/div[2]/span[4]/div/div[3]/div/span').text
            except:
                end = 0
            try:
                T3B4 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[6]/div[2]/span[4]/div/div[3]/div/span').text
            except:
                end = 0
            try:
                T4B4 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[7]/div[2]/span[4]/div/div[3]/div/span').text
            except:
                end = 0
            try:
                T5B4 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[8]/div[2]/span[4]/div/div[3]/div/span').text
            except:
                nope = 0

        T1 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[4]/div[1]').text
        T2 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[5]/div[1]').text
        T3 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[6]/div[1]').text
        T4 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[7]/div[1]').text
        try:
            T5 = driver.find_element_by_xpath('/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/div/div/div[5]/div[2]/div[8]/div[1]').text
        except:
            nope = 0
        exz1 = exz.replace(':','')
        B1 = (int(T1B1)+int(T2B1)+int(T3B1)+int(T4B1)+int(T5B1))/dl
        B2 = (int(T1B2) + int(T2B2) + int(T3B2) + int(T4B2) + int(T5B2)) / dl
        B3 = (int(T1B3) + int(T2B3) + int(T3B3) + int(T4B3) + int(T5B3)) / dl
        if bl == 1:
            B4 = (int(T1B4) + int(T2B4) + int(T3B4) + int(T4B4) + int(T5B4)) / dl
        #questions = (name + ' ' + exz+ '\n'+ '\n' + B1 + '\n' + '\n' + question1 + '\n'+ '\n' + B2 + '\n' +question2 + '\n'+ '\n' + B3 + '\n' +question3 + '\n'+ '\n' + B4 + '\n' + question4)
        doc = DocxTemplate(f"{bl}.docx")
        sc1, sc2, sc3, sc4 = doc.new_subdoc(), doc.new_subdoc(), doc.new_subdoc(), doc.new_subdoc()
        sc1.add_picture(question1, width=Inches(7))
        sc2.add_picture(question2, width=Inches(7))
        sc3.add_picture(question3, width=Inches(7))
        if bl==1:
            sc4.add_picture(question4)
        if bl:
            context = {'B1': B1, 'B2': B2, 'B3': B3, 'B4': B4, 'Fio': name, 'Exz': exz1, 'question1': sc1,
                       'question2': sc2, 'question3': sc3,'T1': T1, 'T2': T2, 'T3': T3, 'T4': T4,
                       'T5': T5, 'T1B1': T1B1, 'T1B2': T1B2, 'T1B3': T1B3, 'T1B4': T1B4, 'T2B1': T2B1, 'T2B2': T2B2,
                       'T2B3': T2B3, 'T2B4': T2B4, 'T3B1': T3B1, 'T3B2': T3B2, 'T3B3': T3B3, 'T3B4': T3B4, 'T4B1': T4B1,
                       'T4B2': T4B2, 'T4B3': T4B3, 'T4B4': T4B4, 'T5B1': T5B1, 'T5B2': T5B2, 'T5B3': T5B3, 'T5B4': T5B4}
        else:
            context = {'B1': B1, 'B2': B2, 'B3': B3, 'B4': B4, 'Fio': name, 'Exz': exz1, 'question1' : sc1, 'question2' : sc2, 'question3' : sc3, 'question4' : sc4, 'T1': T1, 'T2': T2, 'T3': T3, 'T4' : T4, 'T5': T5, 'T1B1': T1B1, 'T1B2': T1B2, 'T1B3' : T1B3, 'T1B4': T1B4, 'T2B1': T2B1, 'T2B2': T2B2, 'T2B3' : T2B3, 'T2B4': T2B4, 'T3B1': T3B1, 'T3B2': T3B2, 'T3B3' : T3B3, 'T3B4': T3B4, 'T4B1': T4B1, 'T4B2': T4B2, 'T4B3' : T4B3, 'T4B4': T4B4, 'T5B1': T5B1, 'T5B2': T5B2, 'T5B3' : T5B3, 'T5B4': T5B4}
        doc.render(context)
        doc.save(f"{name} {exz1}.docx")
        driver.find_element_by_xpath(f'/html/body/div[1]/div/div/div[1]/div[1]/div[2]/div[2]/div[1]/div/div[8]/div/div[2]/div/a[{student_number}]/div[2]/div').click()
        student_number += 1
    driver.get(f'https://app.oqylyq.kz/t/assignments?page=2')
    sleep(1)




